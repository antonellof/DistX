from qdrant_client import models
from qdrant_client import QdrantClient
import os
import openai

# Note: Using qdrant-client library with vectX (Qdrant API-compatible)
# vectX is a fast in-memory vector database: https://github.com/antonellof/vectX

# LLM import - OpenAI only
from llama_index.llms.openai import OpenAI

def batch_iterate(lst, batch_size):
    """Yield successive n-sized chunks from lst."""
    for i in range(0, len(lst), batch_size):
        yield lst[i : i + batch_size]

def extract_text_from_pdf(pdf_path):
    """
    Extract text from PDF using PyMuPDF (fitz).
    This properly handles compressed streams and font encodings.
    """
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(pdf_path)
        text_pages = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                text_pages.append(text)
        
        doc.close()
        
        full_text = "\n\n".join(text_pages)
        print(f"[DEBUG] Extracted {len(full_text)} chars from {len(text_pages)} pages")
        
        # Show first 500 chars for debugging
        if full_text:
            print(f"[DEBUG] Text preview: {full_text[:500]}...")
        
        return full_text
        
    except ImportError:
        print("[DEBUG] PyMuPDF not installed, trying pypdf...")
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(pdf_path)
            text_pages = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_pages.append(text)
            
            full_text = "\n\n".join(text_pages)
            print(f"[DEBUG] Extracted {len(full_text)} chars from {len(text_pages)} pages (pypdf)")
            
            return full_text
            
        except Exception as e:
            print(f"[DEBUG] pypdf also failed: {e}")
            raise
    except Exception as e:
        print(f"[DEBUG] PDF extraction error: {e}")
        raise

def extract_text_from_docx(docx_path):
    """
    Extract text from Word document (.docx).
    """
    try:
        from docx import Document
        
        doc = Document(docx_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        full_text = "\n\n".join(paragraphs)
        print(f"[DEBUG] Extracted {len(full_text)} chars from Word document")
        
        if full_text:
            print(f"[DEBUG] Text preview: {full_text[:500]}...")
        
        return full_text
        
    except ImportError:
        print("[DEBUG] python-docx not installed")
        raise ImportError("python-docx is required for Word documents. Install with: pip install python-docx")
    except Exception as e:
        print(f"[DEBUG] Word extraction error: {e}")
        raise

def extract_text_from_file(file_path):
    """
    Extract text from file based on extension.
    Supports: PDF, DOCX, TXT
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        print(f"[DEBUG] Read {len(text)} chars from text file")
        return text
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def chunk_text(text, chunk_size=1000, overlap=200):
    """Split text into overlapping chunks."""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to break at a sentence or paragraph boundary
        if end < len(text):
            # Look for paragraph break
            last_para = chunk.rfind('\n\n')
            if last_para > chunk_size // 2:
                chunk = chunk[:last_para]
                end = start + last_para
            else:
                # Look for sentence break
                last_period = chunk.rfind('. ')
                if last_period > chunk_size // 2:
                    chunk = chunk[:last_period + 1]
                    end = start + last_period + 1
        
        chunks.append(chunk.strip())
        start = end - overlap  # Overlap for context continuity
        
        if start >= len(text):
            break
    
    return [c for c in chunks if c]  # Remove empty chunks

class EmbedData:
    """
    Embedding class using OpenAI embeddings.
    Fast startup - no local models to load!
    Includes automatic text chunking for large documents.
    """

    def __init__(self, embed_model_name="text-embedding-3-small", batch_size=20, chunk_size=1000):
        self.embed_model_name = embed_model_name
        self.batch_size = batch_size  # Number of chunks per API call
        self.chunk_size = chunk_size  # Characters per chunk
        self.embeddings = []
        self.contexts = []
        
        # Initialize OpenAI client
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable is required")
        self.client = openai.OpenAI(api_key=api_key)
        
        # Dimension for text-embedding-3-small is 1536
        self.dimension = 1536

    def generate_embedding(self, texts):
        """Generate embeddings using OpenAI API."""
        # Filter out empty texts
        valid_texts = [t for t in texts if t and t.strip()]
        if not valid_texts:
            return []
            
        response = self.client.embeddings.create(
            model=self.embed_model_name,
            input=valid_texts
        )
        return [item.embedding for item in response.data]
        
    def embed(self, documents):
        """Chunk and embed all documents."""
        # First, chunk all documents
        all_chunks = []
        for doc in documents:
            chunks = chunk_text(doc, chunk_size=self.chunk_size)
            all_chunks.extend(chunks)
        
        self.contexts = all_chunks
        
        # Embed in batches
        for batch_context in batch_iterate(all_chunks, self.batch_size):
            batch_embeddings = self.generate_embedding(batch_context)
            self.embeddings.extend(batch_embeddings)
    
    def get_query_embedding(self, query):
        """Get embedding for a single query."""
        response = self.client.embeddings.create(
            model=self.embed_model_name,
            input=[query]
        )
        return response.data[0].embedding

class vectXVDB:
    """
    vectX Vector Database client.
    Uses qdrant-client library for Qdrant API compatibility.
    vectX is a fast in-memory vector database: https://github.com/antonellof/vectX
    """

    def __init__(self, collection_name, vector_dim = 768, batch_size=512):
        self.collection_name = collection_name
        self.batch_size = batch_size
        self.vector_dim = vector_dim
        
    def define_client(self):
        # vectX uses the same API as Qdrant, so qdrant-client works seamlessly
        # prefer_grpc=False since vectX gRPC might have different implementation
        # check_compatibility=False to skip version check (vectX doesn't report version like Qdrant)
        self.client = QdrantClient(
            url="http://localhost:6333", 
            prefer_grpc=False,
            check_compatibility=False
        )
        
    def create_collection(self):
        """Create a collection in vectX. Uses simplified config compatible with vectX."""
        if not self.client.collection_exists(collection_name=self.collection_name):
            # vectX supports basic vector config: size and distance
            # Removed: on_disk, optimizers, quantization (not yet supported in vectX)
            self.client.create_collection(
                collection_name=f"{self.collection_name}",
                vectors_config=models.VectorParams(
                    size=self.vector_dim,
                    distance=models.Distance.COSINE,  # Changed to COSINE (better for embeddings)
                )
            )
            
    def ingest_data(self, embeddata, document_name="unknown"):
        """Ingest embeddings into vectX collection with document tracking."""
        import time
        import uuid
        
        # Get current max ID to avoid conflicts
        try:
            # Use UUID-based IDs to avoid conflicts when adding new documents
            base_id = int(time.time() * 1000) % 1000000000
        except:
            base_id = 0
        
        # Convert embeddings to PointStruct format for qdrant-client
        points = []
        for i, (context, embedding) in enumerate(zip(embeddata.contexts, embeddata.embeddings)):
            points.append(models.PointStruct(
                id=base_id + i,
                vector=embedding,
                payload={
                    "context": context,
                    "document": document_name,  # Track which document this chunk came from
                    "chunk_index": i
                }
            ))
        
        print(f"[DEBUG] Ingesting {len(points)} points from '{document_name}' into '{self.collection_name}'")
        print(f"[DEBUG] Vector dimension: {len(embeddata.embeddings[0]) if embeddata.embeddings else 'N/A'}")
        
        # Batch upsert points
        batch_count = 0
        for batch_points in batch_iterate(points, self.batch_size):
            self.client.upsert(
                collection_name=self.collection_name,
                points=batch_points
            )
            batch_count += 1
        
        print(f"[DEBUG] Ingested {len(points)} points in {batch_count} batches")
        
        # Verify ingestion
        try:
            info = self.client.get_collection(self.collection_name)
            print(f"[DEBUG] Collection now has {info.points_count} points")
        except Exception as e:
            print(f"[DEBUG] Could not verify ingestion: {e}")
    
    def list_documents(self):
        """List unique documents in the collection with their chunk counts."""
        try:
            # Scroll through all points to get unique document names
            documents = {}
            offset = None
            
            while True:
                results, offset = self.client.scroll(
                    collection_name=self.collection_name,
                    limit=100,
                    offset=offset,
                    with_payload=True,
                    with_vectors=False
                )
                
                for point in results:
                    if point.payload and "document" in point.payload:
                        doc_name = point.payload["document"]
                        if doc_name not in documents:
                            documents[doc_name] = {"name": doc_name, "chunks": 0, "point_ids": []}
                        documents[doc_name]["chunks"] += 1
                        documents[doc_name]["point_ids"].append(point.id)
                
                if offset is None:
                    break
            
            return list(documents.values())
        except Exception as e:
            print(f"[DEBUG] Error listing documents: {e}")
            return []
    
    def delete_document(self, document_name):
        """Delete all points belonging to a specific document."""
        try:
            # Use filter to delete points by document name
            self.client.delete(
                collection_name=self.collection_name,
                points_selector=models.FilterSelector(
                    filter=models.Filter(
                        must=[
                            models.FieldCondition(
                                key="document",
                                match=models.MatchValue(value=document_name)
                            )
                        ]
                    )
                )
            )
            print(f"[DEBUG] Deleted all points from document '{document_name}'")
            return True
        except Exception as e:
            print(f"[DEBUG] Error deleting document: {e}")
            return False
        
class Retriever:

    def __init__(self, vector_db, embeddata):
        
        self.vector_db = vector_db
        self.embeddata = embeddata

    def search(self, query):
        """Search for similar vectors in vectX."""
        query_embedding = self.embeddata.get_query_embedding(query)
        
        # Debug: check collection info
        try:
            collection_info = self.vector_db.client.get_collection(self.vector_db.collection_name)
            print(f"[DEBUG] Collection '{self.vector_db.collection_name}' has {collection_info.points_count} points")
        except Exception as e:
            print(f"[DEBUG] Error getting collection info: {e}")
        
        # Search
        result = self.vector_db.client.search(
            collection_name=self.vector_db.collection_name,
            query_vector=query_embedding,
            limit=10,
        )
        
        print(f"[DEBUG] Search returned {len(result)} results")
        if result:
            print(f"[DEBUG] First result: {result[0]}")

        return result
    
class RAG:

    def __init__(self,
                 retriever,
                 llm_name="gpt-5-mini",   # Model name (default: GPT-5 mini)
                 api_key=None             # OpenAI API key
                 ):
        
        self.llm_name = llm_name
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        self.llm = self._setup_llm()
        self.retriever = retriever
        self.qa_prompt_tmpl_str = ("Context information is below.\n"
                                   "---------------------\n"
                                   "{context}\n"
                                   "---------------------\n"
                                   "Given the context information above I want you to think step by step to answer the query in a crisp manner, incase case you don't know the answer say 'I don't know!'.\n"
                                   "Query: {query}\n"
                                   "Answer: "
                                   )

    def _setup_llm(self):
        """Setup OpenAI LLM."""
        if not self.api_key:
            raise ValueError("OPENAI_API_KEY environment variable is required. Set it with: export OPENAI_API_KEY=your-key")
        
        return OpenAI(
            model=self.llm_name,
            temperature=0.7,
            api_key=self.api_key,
        )

    def generate_context(self, query):

        result = self.retriever.search(query)
        combined_prompt = []
        
        print(f"[DEBUG] Processing {len(result)} search results for context")

        for i, scored_point in enumerate(result[:3]):  # Use top 3 results
            try:
                # ScoredPoint has .payload attribute (not dict access)
                if hasattr(scored_point, 'payload') and scored_point.payload:
                    payload = scored_point.payload
                    if isinstance(payload, dict) and "context" in payload:
                        context_text = payload["context"]
                        combined_prompt.append(context_text)
                        print(f"[DEBUG] Result {i}: score={scored_point.score:.4f}, context_len={len(context_text)}")
                    else:
                        print(f"[DEBUG] Result {i}: payload has no 'context' key: {payload.keys() if isinstance(payload, dict) else type(payload)}")
                else:
                    print(f"[DEBUG] Result {i}: no payload attribute")
            except Exception as e:
                print(f"[DEBUG] Error processing result {i}: {e}")

        if not combined_prompt:
            print("[DEBUG] WARNING: No context extracted from search results!")
            return "No relevant context found in the document."
            
        return "\n\n---\n\n".join(combined_prompt)

    def query(self, query):
        """Query the RAG system and return streaming response."""
        context = self.generate_context(query=query)
        
        prompt = self.qa_prompt_tmpl_str.format(context=context, query=query)
                
        # OpenAI supports stream_complete
        streaming_response = self.llm.stream_complete(prompt)
        
        return streaming_response
    
    # def append_ai_response(self, message):

    #     self.messages.append(ChatMessage(role=MessageRole.ASSISTANT, content=message))