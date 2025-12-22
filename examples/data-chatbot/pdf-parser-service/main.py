"""
FastAPI service for parsing PDF, Word, and text documents.
Uses the same approach as fastest-rag-stack for PDF extraction.
"""

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import io
import os

app = FastAPI(title="PDF Parser Service")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, restrict this
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ParseResponse(BaseModel):
    text: str
    pageCount: Optional[int] = None
    wordCount: int


def extract_text_from_pdf(file_buffer: bytes) -> tuple[str, int]:
    """
    Extract text from PDF using PyMuPDF (fitz).
    This properly handles compressed streams and font encodings.
    Returns: (text, page_count)
    """
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(stream=file_buffer, filetype="pdf")
        text_pages = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                text_pages.append(text)
        
        doc.close()
        
        full_text = "\n\n".join(text_pages)
        return full_text, len(text_pages)
        
    except ImportError:
        # Fallback to pypdf
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(io.BytesIO(file_buffer))
            text_pages = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_pages.append(text)
            
            full_text = "\n\n".join(text_pages)
            return full_text, len(text_pages)
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF extraction failed: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF extraction error: {str(e)}")


def extract_text_from_docx(file_buffer: bytes) -> str:
    """
    Extract text from Word document (.docx).
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_buffer))
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        full_text = "\n\n".join(paragraphs)
        return full_text
        
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is required for Word documents. Install with: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word extraction error: {str(e)}")


def extract_text_from_txt(file_buffer: bytes) -> str:
    """
    Extract text from plain text file.
    """
    try:
        text = file_buffer.decode("utf-8", errors="ignore")
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Text file reading error: {str(e)}")


@app.get("/health")
async def health():
    """Health check endpoint."""
    return {"status": "ok"}


@app.post("/parse", response_model=ParseResponse)
async def parse_document(file: UploadFile = File(...)):
    """
    Parse a document (PDF, Word, or text) and extract text.
    """
    # Read file content
    file_buffer = await file.read()
    
    if len(file_buffer) == 0:
        raise HTTPException(status_code=400, detail="File is empty")
    
    # Get file extension
    filename = file.filename or "unknown"
    ext = os.path.splitext(filename)[1].lower()
    
    try:
        if ext == ".pdf":
            text, page_count = extract_text_from_pdf(file_buffer)
            word_count = len(text.split())
            return ParseResponse(
                text=text,
                pageCount=page_count,
                wordCount=word_count
            )
        elif ext in [".docx", ".doc"]:
            text = extract_text_from_docx(file_buffer)
            word_count = len(text.split())
            return ParseResponse(
                text=text,
                pageCount=None,  # Word docs don't have a reliable page count
                wordCount=word_count
            )
        elif ext in [".txt", ".text"]:
            text = extract_text_from_txt(file_buffer)
            word_count = len(text.split())
            return ParseResponse(
                text=text,
                pageCount=None,
                wordCount=word_count
            )
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT"
            )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8000"))
    uvicorn.run(app, host="0.0.0.0", port=port)
